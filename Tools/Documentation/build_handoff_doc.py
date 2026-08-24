from __future__ import annotations

import argparse
import importlib.util
import json
import math
import re
import subprocess
from collections import Counter
from datetime import datetime
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "Documentation"
QA_DIR = OUT_DIR / "_qa_handoff"
IMAGE_DIR = QA_DIR / "images"
OUT_DOCX = OUT_DIR / "CraftOrigin_Craft-live_アカウント引継ぎ詳細書_2026-08-09.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F8FC"
RED = "C00000"
PALE_RED = "FCE8E6"
AMBER = "B36B00"
PALE_AMBER = "FFF4CE"
GREEN = "2E7D32"
PALE_GREEN = "EAF5EA"
GRAY = "666666"
LIGHT_GRAY = "F2F2F2"
WHITE = "FFFFFF"
BLACK = "1F1F1F"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_heading(table) -> None:
    if table.rows:
        repeat_table_header(table.rows[0])


def set_run_font(run, name="Yu Gothic", size=None, bold=None, color=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.15, keep_next=False) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_next


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Yu Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    heading_specs = {
        "Title": (26, DARK_BLUE, 0, 12),
        "Subtitle": (13, GRAY, 0, 10),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (11.5, DARK_BLUE, 10, 5),
        "Heading 4": (10.2, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = "Yu Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Yu Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
        style.font.size = Pt(9.2)
        style.paragraph_format.space_after = Pt(2.5)
        style.paragraph_format.left_indent = Inches(0.25)
        style.paragraph_format.first_line_indent = Inches(-0.15)

    if "Table Text" not in styles:
        table_text = styles.add_style("Table Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        table_text = styles["Table Text"]
    table_text.font.name = "Yu Gothic"
    table_text._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    table_text.font.size = Pt(7.7)
    table_text.paragraph_format.space_after = Pt(0)
    table_text.paragraph_format.line_spacing = 1.05

    if "Small Note" not in styles:
        small = styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        small = styles["Small Note"]
    small.font.name = "Yu Gothic"
    small._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    small.font.size = Pt(7.8)
    small.font.color.rgb = RGBColor.from_string(GRAY)
    small.paragraph_format.space_after = Pt(3)
    small.paragraph_format.line_spacing = 1.08

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("CRAFTORIGIN / CRAFT-LIVE  |  HANDOFF REFERENCE")
    set_run_font(r, size=7.5, bold=True, color=GRAY)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), LIGHT_BLUE)
    pbdr.append(bottom)
    pPr.append(pbdr)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Confidential / account handoff reference   |   ")
    set_run_font(r, size=7.2, color=GRAY)
    add_field(p, "PAGE")
    r = p.add_run(" / ")
    set_run_font(r, size=7.2, color=GRAY)
    add_field(p, "NUMPAGES")


def set_table_borders(table, color="D9E2F3", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def add_table(doc, headers, rows, widths, font_size=7.7, header_fill=LIGHT_BLUE, zebra=True):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_widths(table, widths)
    set_table_borders(table)
    table.autofit = False
    header = table.rows[0]
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        shade_cell(cell, header_fill)
        p = cell.paragraphs[0]
        p.style = "Table Text"
        r = p.add_run(str(text))
        set_run_font(r, size=font_size, bold=True, color=DARK_BLUE)
    repeat_table_header(header)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        if zebra and ridx % 2 == 1:
            for cell in cells:
                shade_cell(cell, "F8FAFD")
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.style = "Table Text"
            text = "" if value is None else str(value)
            r = p.add_run(text)
            set_run_font(r, size=font_size)
    return table


def add_callout(doc, title, body, level="info"):
    colors = {
        "critical": (RED, PALE_RED),
        "warning": (AMBER, PALE_AMBER),
        "success": (GREEN, PALE_GREEN),
        "info": (BLUE, PALE_BLUE),
    }
    accent, fill = colors[level]
    table = doc.add_table(rows=1, cols=2)
    set_table_widths(table, [140, 9220])
    set_table_borders(table, color=fill, size="0")
    # Treat the single semantic callout row as its accessible header row.
    repeat_table_header(table.rows[0])
    shade_cell(table.cell(0, 0), accent)
    shade_cell(table.cell(0, 1), fill)
    p = table.cell(0, 1).paragraphs[0]
    p.style = "Table Text"
    r = p.add_run(title + "\n")
    set_run_font(r, size=9, bold=True, color=accent)
    r = p.add_run(body)
    set_run_font(r, size=8.3, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_bullets(doc, items, numbered=False, level=0):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        if level:
            p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
        if isinstance(item, tuple):
            lead, rest = item
            r = p.add_run(lead)
            set_run_font(r, bold=True)
            r = p.add_run(rest)
            set_run_font(r)
        else:
            r = p.add_run(str(item))
            set_run_font(r)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style = "Small Note"
    r = p.add_run(text)
    set_run_font(r, size=7.6, color=GRAY)


def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def get_font(size: int, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/YuGothB.ttc" if bold else "C:/Windows/Fonts/YuGothR.ttc"),
        Path("C:/Windows/Fonts/meiryob.ttc" if bold else "C:/Windows/Fonts/meiryo.ttc"),
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size, index=0)
    return ImageFont.load_default()


def draw_round_box(draw, box, fill, outline, title, lines, title_color=DARK_BLUE):
    x0, y0, x1, y1 = box
    draw.rounded_rectangle(box, radius=18, fill="#" + fill, outline="#" + outline, width=3)
    title_font = get_font(32, bold=True)
    body_font = get_font(22)
    draw.text((x0 + 24, y0 + 18), title, font=title_font, fill="#" + title_color)
    y = y0 + 70
    for line in lines:
        draw.text((x0 + 26, y), line, font=body_font, fill="#" + BLACK)
        y += 34


def make_architecture_diagram(path: Path):
    img = Image.new("RGB", (1800, 1080), "white")
    d = ImageDraw.Draw(img)
    title = get_font(42, bold=True)
    d.text((70, 42), "Craft-live 現行アーキテクチャ", font=title, fill="#" + DARK_BLUE)
    pad_boxes = [
        ((80, 160, 760, 330), "Pad 1 / Material Gallery", ["素材閲覧・常設登録", "配置候補の選択・発射キュー"]),
        ((80, 370, 760, 540), "Pad 2 / Workbench", ["武器選択・6スロット配置", "液体演出・ハンマー合成・結果"]),
        ((80, 580, 760, 750), "Pad 3 / Status + QR", ["QR読取・素材登録", "攻撃 / 防御 / 回避ゲージ表示"]),
        ((80, 790, 760, 960), "Pad 4 / Hologram", ["完成武器ホログラム", "最終コード・履歴表示"]),
    ]
    for box, t, lines in pad_boxes:
        draw_round_box(d, box, PALE_BLUE, BLUE, t, lines)
    draw_round_box(d, (1030, 205, 1710, 440), LIGHT_BLUE, BLUE, "Firebase Realtime Database", ["/rooms/{roomId}.json", "RoomState schema v5", "REST + ETag / conditional PUT"])
    draw_round_box(d, (1030, 560, 1710, 820), PALE_GREEN, GREEN, "Bootstrap + 共通ランタイム", ["URL: screen / pad / room", "Session / Transport / Timer", "WebPresentation / Diagnostics", "役割別SceneをAdditive Load"], title_color=GREEN)
    for y in (245, 455, 665, 875):
        d.line((760, y, 1030, 325), fill="#" + BLUE, width=5)
        d.polygon([(1030, 325), (1005, 312), (1005, 338)], fill="#" + BLUE)
    d.line((1370, 440, 1370, 560), fill="#" + GREEN, width=6)
    d.polygon([(1370, 560), (1356, 532), (1384, 532)], fill="#" + GREEN)
    note = get_font(20)
    d.text((1040, 870), "4台は同じ roomId を共有し、状態を同期する。", font=note, fill="#" + GRAY)
    img.save(path)


def make_flow_diagram(path: Path):
    img = Image.new("RGB", (1800, 950), "white")
    d = ImageDraw.Draw(img)
    title = get_font(42, bold=True)
    d.text((70, 42), "体験フローと状態遷移", font=title, fill="#" + DARK_BLUE)
    steps = [
        ("1", "ルーム開始", "4台を同じroomで起動"),
        ("2", "素材登録", "Pad3でQR / QR不要素材"),
        ("3", "武器選択", "Pad2で武器を確定"),
        ("4", "素材転送", "Pad1→キュー→Pad2到着"),
        ("5", "スロット配置", "属性・スキル・4基礎枠"),
        ("6", "合成", "液体→6打撃→ランク"),
        ("7", "結果共有", "Pad3ステータス / Pad4表示"),
        ("8", "最終決定", "履歴から選びCLコード生成"),
    ]
    x_positions = [90, 505, 920, 1335]
    y_positions = [170, 500]
    idx = 0
    for row, y in enumerate(y_positions):
        items = steps[row * 4 : row * 4 + 4]
        if row == 1:
            items = list(reversed(items))
        for col, item in enumerate(items):
            x = x_positions[col]
            num, title_text, detail = item
            fill = PALE_GREEN if num in ("6", "7", "8") else PALE_BLUE
            outline = GREEN if num in ("6", "7", "8") else BLUE
            draw_round_box(d, (x, y, x + 335, y + 200), fill, outline, f"{num}. {title_text}", [detail], title_color=outline)
        if row == 0:
            for x in (425, 840, 1255):
                d.line((x, y + 100, x + 80, y + 100), fill="#" + BLUE, width=5)
                d.polygon([(x + 80, y + 100), (x + 55, y + 86), (x + 55, y + 114)], fill="#" + BLUE)
        else:
            for x in (505, 920, 1335):
                d.line((x, y + 100, x - 80, y + 100), fill="#" + GREEN, width=5)
                d.polygon([(x - 80, y + 100), (x - 55, y + 86), (x - 55, y + 114)], fill="#" + GREEN)
    d.line((1670, 370, 1670, 500), fill="#" + BLUE, width=5)
    d.polygon([(1670, 500), (1656, 472), (1684, 472)], fill="#" + BLUE)
    note = get_font(22)
    d.text((90, 785), "状態は revision / updatedAt で比較。セッションは300秒、履歴は最大12件。", font=note, fill="#" + GRAY)
    d.text((90, 830), "現行Rulesでは属性・スキル・4基礎枠はすべて必須ではないため、武器確定後は素材0件でも合成開始できる。", font=note, fill="#" + RED)
    img.save(path)


def make_ownership_diagram(path: Path):
    img = Image.new("RGB", (1800, 900), "white")
    d = ImageDraw.Draw(img)
    title = get_font(42, bold=True)
    d.text((70, 42), "引継ぎ境界：ファイルで渡るもの / アカウント操作が必要なもの", font=title, fill="#" + DARK_BLUE)
    draw_round_box(d, (80, 145, 830, 765), PALE_BLUE, BLUE, "プロジェクト内で渡る", [
        "Assets / Packages / ProjectSettings",
        "C#コード・Scene・Prefab・素材定義",
        "WebGLテンプレート・QRブリッジ",
        "FirebaseのURL・ルーム形式",
        "Git設定・LFS属性（ただし現状は未追跡多数）",
        "この引継ぎ詳細書",
    ])
    draw_round_box(d, (970, 145, 1720, 765), PALE_AMBER, AMBER, "アカウント側で別途移管", [
        "GitHub repository owner / collaborator",
        "Git LFSの実体と容量・課金",
        "Firebase / Google Cloudの所有権・Rules・データ",
        "公開先ホスティング・独自ドメイン・TLS",
        "Unity ID / ライセンス / Unity Cloud権限",
        "Meshyアカウント・API権限",
        "端末・Wi-Fi・MDM・運用手順",
    ], title_color=AMBER)
    d.line((830, 455, 970, 455), fill="#" + RED, width=7)
    d.polygon([(970, 455), (936, 438), (936, 472)], fill="#" + RED)
    f = get_font(22, bold=True)
    d.text((710, 805), "両方が揃って初めて復旧可能", font=f, fill="#" + RED)
    img.save(path)


def add_picture_with_alt(doc, path: Path, width=6.8, alt_text=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    inline_shape = run.add_picture(str(path), width=Inches(width))
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", alt_text)
    return inline_shape


def load_scene_analysis():
    module_path = ROOT / "Tools" / "Documentation" / "analyze_project.py"
    spec = importlib.util.spec_from_file_location("analyze_project", module_path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    guid_map = module.build_guid_map()
    scenes = []
    for path in sorted((ROOT / "Assets" / "Scenes").rglob("*.unity")):
        scenes.append(module.parse_scene(path, guid_map))
    return scenes


def collect_scripts():
    files = sorted(list((ROOT / "Assets" / "Scripts" / "CraftLive").rglob("*.cs")) + list((ROOT / "Assets" / "Editor").glob("CraftLive*.cs")))
    result = []
    decl_re = re.compile(r"(?m)^\s*(?:public|internal)\s+(?:(?:sealed|static|abstract|partial)\s+)*(class|struct|enum|interface)\s+(\w+)")
    field_re = re.compile(r"\[SerializeField(?:[^\]]*)\]\s*(?:(?:\[[^\]]+\])\s*)*(?:private|protected)\s+([\w<>,\[\].?]+)\s+(\w+)", re.S)
    for path in files:
        text = path.read_text(encoding="utf-8", errors="replace")
        decls = [f"{kind} {name}" for kind, name in decl_re.findall(text)]
        fields = [(typ, name) for typ, name in field_re.findall(text)]
        result.append({"path": path.relative_to(ROOT).as_posix(), "name": path.stem, "decls": decls, "fields": fields, "lines": text.count("\n") + 1})
    return result


SCRIPT_DESCRIPTIONS = {
    "CraftLiveTypes": "共有Enum、Stats、属性/スキル効果、スロット補助。JSON/Sceneの値契約を定義する。",
    "CraftLiveRoomState": "RoomState schema v5、転送・配置・合成・結果・履歴・最終選択を保持し、旧schemaをNormalizeする。",
    "CraftLiveCalculator": "武器基礎値、4基礎スロット補正、合成ランクボーナスから最終ステータスを計算する。",
    "CraftLiveWeaponCode": "武器/属性/スキル/攻撃・防御・回避素材数を6文字へ直接符号化し、オフラインで復号する。",
    "CraftLiveCatalog": "素材と武器のScriptableObject一覧。ID検索と定義参照の入口。",
    "CraftLiveMaterialDefinition": "素材ID、表示、カテゴリ、QR要否、Prefab/Icon、能力値・演出参照を定義する。",
    "CraftLiveWeaponDefinition": "武器ID、名称、種類、基礎能力、作業台/ホログラムPrefab、プレビュー値を定義する。",
    "CraftLiveRules": "セッション時間、必須枠、ハンマー回数、ランク閾値、履歴上限、最大能力値を定義する。",
    "CraftLiveLaunchConfig": "役割Scene、room、Firebase URL、poll/timeout/retry/cacheを定義する。",
    "CraftLivePad4Calibration": "Pad4表示面の物理寸法、角度、距離、位置/回転/scaleを定義する。",
    "CraftLiveSession": "状態変更APIの中心。登録、選択、転送、配置、合成、履歴、最終決定を一貫して更新する。",
    "CraftLiveRoomTransport": "Editorローカル共有またはFirebase REST同期、ETag条件付きPUT、再試行、PlayerPrefsキャッシュを担当する。",
    "CraftLiveBootstrap": "URLクエリからrole/roomを解決し、共通参照を初期化、役割SceneをAdditive Loadする。",
    "CraftLiveRoleRouter": "起動roleに対応するPad Sceneをロードする。",
    "CraftLiveLaunchQuery": "screen/pad/roomクエリの抽出と別名解決を行う。",
    "CraftLivePlacementWatchdog": "Pad1/Pad2転送状態が長時間停滞した場合に段階を進める復旧補助。",
    "CraftLiveQrScanner": "WebGL JavaScript QRオーバーレイの開始・終了・結果コールバックをUnity側へ橋渡しする。",
    "CraftLiveSessionTimerController": "300秒タイマーとFinalSelectionへの遷移を制御する。",
    "CraftLiveRuntimeDiagnostics": "同期・Scene・状態の異常やstale状態を定期確認してログ化する。",
    "CraftLiveWebPresentation": "30fps、vSync、スリープ禁止、Safe Area、3:4表示を統一する。",
    "CraftLivePadSceneRoot": "各Pad Sceneのroleを示す識別コンポーネント。",
    "CraftLiveCommandActions": "リセットや進行など、UIからSessionのコマンドを呼ぶ。",
    "CraftLiveMaterialAction": "素材選択/登録等のUIアクションをSessionへ接続する。",
    "CraftLiveWeaponAction": "武器選択/確定のUIアクションをSessionへ接続する。",
    "CraftLiveSlotAction": "スロット候補・確定・解除のUIアクションをSessionへ接続する。",
    "CraftLiveStateEvents": "状態更新をUnityEventへ変換してScene上の表示更新を疎結合にする。",
    "CraftLiveWorldButton": "3D空間のクリック/タップ可能ボタン。",
    "CraftLiveMaterialTicketView": "素材チケット表示と入力。",
    "CraftLiveMaterialBoardView": "素材一覧ボード表示。",
    "CraftLivePlacementSlotView": "Pad2の配置枠表示と状態色/入力。",
    "CraftLiveMixInput": "旧/補助の回転ミックス入力。現行主合成はHammer controller。",
    "CraftLivePad1Bindings": "Pad1 Scene参照をまとめるバインディング。",
    "CraftLivePad1GalleryController": "素材定義から画廊を構成し、ロック/選択/スクロールを管理する。",
    "CraftLivePad1MaterialPreview": "素材3Dプレビュー、説明パネル、転送/戻る操作、選択状態を管理する。",
    "CraftLivePad1TransferController": "バネ引き、転送キュー、装填/発射/到着演出とRoomState更新を管理する。",
    "CraftLiveGalleryColumn": "1列分の作品・入力面を生成/配置する。",
    "CraftLiveGalleryWallSlider": "壁面のスワイプ/スナップ移動と表示範囲を管理する。",
    "CraftLiveGalleryWallView": "画廊壁面全体の表示構造。",
    "CraftLiveMaterialPaintingView": "額装された素材作品とロック/選択表示。",
    "CraftLivePad1Presentation": "Pad1表示用の色・レイアウト補助を提供する。",
    "CraftLiveTransferLauncherView": "発射装置の視覚演出。",
    "CraftLivePad2Bindings": "Pad2 Scene参照をまとめるバインディング。",
    "CraftLivePad2WeaponCarousel": "8武器の横カルーセル、スワイプ、選択/確定、プレビュー生成。",
    "CraftLivePad2PlacementController": "6スロットの候補選択、確定、解除、重複防止。",
    "CraftLivePad2TransferReceiver": "Pad1キューを受け、到着軌道、溝移動、配置完了を演出する。",
    "CraftLiveLiquidFlowController": "配置後の液体滴下/プール演出。",
    "CraftLiveHammerSynthesisController": "ハンマー入力6回、強さ計算、ランク確定、Session合成完了を制御する。",
    "CraftLivePad2ResultController": "完成結果、ランク、能力値、次の武器/最終選択を表示する。",
    "CraftLivePad2SlotLayout": "論理スロットと物理位置の対応を固定する。",
    "CraftLivePad2AlignmentGuide": "作業台上の位置合わせ用ガイド。",
    "CraftLiveWorkbenchView": "作業台の武器/素材表示。",
    "CraftLivePad3Bindings": "Pad3 Scene参照をまとめるバインディング。",
    "CraftLivePad3Controller": "QR操作、登録メッセージ、3本のステータス管を状態に同期する。",
    "CraftLiveStatusTubeView": "攻撃/防御/回避の液面アニメーション。",
    "CraftLivePad4Bindings": "Pad4 Scene参照をまとめるバインディング。",
    "CraftLivePad4Controller": "完成履歴、選択結果、最終コード、ホログラム表示を同期する。",
    "CraftLiveHologramView": "武器Prefabの生成、回転、属性色/発光、キャリブレーション適用。",
    "CraftLiveGeneratedRuntimeVisual": "実行時自動生成Visualを識別し再生成時に整理するマーカー。",
    "CraftLiveDefaultDataCreator": "既定Catalog/Rules/LaunchConfig/Calibration/定義Assetの作成・更新。",
    "CraftLiveMaterialDefinitionEditor": "素材Definition Inspectorの編集支援と不備警告。",
    "CraftLiveStep0BaselineValidator": "現行データ、Build Scene、各Scene参照を横断検証しMarkdownレポートを生成。",
    "CraftLiveStep9ProductionUpgrader": "WebGL本番設定の適用・検証・ビルド支援。",
    "CraftLiveEditModeTestRunner": "EditModeテストの一括実行補助。",
}


def infer_script_description(name: str, path: str) -> str:
    if name in SCRIPT_DESCRIPTIONS:
        return SCRIPT_DESCRIPTIONS[name]
    if name.endswith("Tests") or "Tests" in name:
        return "対応Step/機能のEditMode回帰テスト。旧実装前提が残る場合は現行仕様と照合して保守する。"
    if name.endswith("Upgrader") or "Setup" in name or "Generator" in name:
        return "開発段階のScene/Asset生成・移行用Editorツール。再実行前に現行Sceneとの差分を確認する。"
    return "Craft-live実装コンポーネント。詳細な[SerializeField]は付録Dを参照。"


def collect_manifest():
    rows = []
    for base_name in ("Assets", "Packages", "ProjectSettings"):
        base = ROOT / base_name
        for path in sorted(p for p in base.rglob("*") if p.is_file()):
            rel = path.relative_to(ROOT).as_posix()
            ext = path.suffix.lower() or "(none)"
            size = path.stat().st_size
            if ext == ".meta":
                role = "Unity GUID/import metadata（必須）"
            elif ext == ".cs":
                role = "C# source / Editor test or tool"
            elif ext == ".unity":
                role = "Unity Scene"
            elif ext == ".prefab":
                role = "Prefab"
            elif ext == ".asset":
                role = "ScriptableObject / render settings"
            elif ext in (".obj", ".fbx"):
                role = "3D mesh source"
            elif ext in (".png", ".jpg", ".jpeg", ".tga"):
                role = "Texture / icon"
            elif ext == ".mat":
                role = "Material"
            elif ext == ".md":
                role = "Project documentation"
            elif ext in (".html", ".css", ".jslib"):
                role = "WebGL template / browser bridge"
            elif ext == ".json":
                role = "Package/configuration manifest"
            elif ext == ".asmdef":
                role = "Assembly definition"
            else:
                role = "Project configuration / supporting asset"
            rows.append((rel, ext, size, role))
    return rows


def fmt_size(size: int) -> str:
    if size >= 1024 * 1024:
        return f"{size / 1024 / 1024:.2f} MB"
    if size >= 1024:
        return f"{size / 1024:.1f} KB"
    return f"{size} B"


def add_cover(doc):
    p = doc.add_paragraph()
    r = p.add_run("PROJECT HANDOFF  /  UNITY WEBGL")
    set_run_font(r, size=9, bold=True, color=BLUE)
    p.paragraph_format.space_after = Pt(54)

    p = doc.add_paragraph()
    p.style = "Title"
    r = p.add_run("CraftOrigin / Craft-live")
    set_run_font(r, size=28, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.style = "Subtitle"
    r = p.add_run("アカウント引継ぎ・プロジェクト全体詳細書")
    set_run_font(r, size=16, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(34)
    r = p.add_run("現行ワークスペース基準 / 実装・設定・運用・外部サービスを統合")
    set_run_font(r, size=10.5, color=GRAY)

    add_callout(
        doc,
        "最優先の引継ぎ警告",
        "Gitの追跡対象は66ファイルだけで、Craft-live本体を含む527ファイルが未追跡です（本書生成物を除く）。GitHubからcloneするだけでは現行プロジェクトを復元できません。外部アカウント移管より先に、完全なAssets/Packages/ProjectSettingsをGit LFS対応で保全してください。",
        "critical",
    )

    meta = [
        ("基準日", "2026-08-09 JST"),
        ("Unity", "6000.4.0f1 / URP 17.4.0 / WebGL"),
        ("製品名", "Craft-live（Company: goodmorning2424 / Version: 0.1.0）"),
        ("対象", "4台iPad連携クラフト体験 / Firebase Realtime Database"),
        ("文書種別", "復旧可能性を優先した、現行実装ベースの詳細リファレンス"),
    ]
    add_table(doc, ["項目", "内容"], meta, [1900, 7460], font_size=8.4)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    r = p.add_run("読み方")
    set_run_font(r, size=10, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph("本文の「確認済み」は現行Scene/Asset/コードまたは2026-08-08の自動検証で確認した事実です。「履歴」は旧Step文書にのみ残る内容、「要移管」は外部アカウントでの操作が必要な内容です。")
    p.style = "Small Note"


def add_document_control(doc):
    add_page_break(doc)
    doc.add_heading("文書管理と情報の優先順位", level=1)
    add_table(doc, ["優先", "情報源", "扱い"], [
        ("1", "現行コード / .asset / .unity / ProjectSettings", "実行時の真実。本文の仕様判断は原則これを採用。"),
        ("2", "2026-08-08 CurrentValidation / ProductionReadiness / EditMode test", "自動検証で再確認した状態。"),
        ("3", "CRAFTLIVE_PROJECT_OVERVIEW_JA / INSPECTOR GUIDE", "全体説明として有用だが、Catalog件数や一部設定は現行との差分あり。"),
        ("4", "STEP0〜STEP9文書 / PROJECT_STATE_JA", "設計変更と構築履歴。現行値と衝突する場合は履歴扱い。"),
        ("5", "README_SETUP_JA", "V2時代の旧方式。消費在庫・旧Pad1 QR・11素材/3武器の記述は現行仕様ではない。"),
    ], [700, 3300, 5360], font_size=8.1)
    doc.add_heading("本書が扱う範囲", level=2)
    add_bullets(doc, [
        "アカウント移管：GitHub/Git LFS、Unity ID、Firebase/Google Cloud、公開ホスティング、Meshy、端末運用。",
        "プロジェクト全体像：4画面、状態モデル、同期、データ定義、計算、WebGL/QR、Scene構造。",
        "完全性確認：現行Catalog、Rules、Build Settings、パッケージ、スクリプト、全Assets/設定ファイルの一覧。",
        "復旧・継続開発：取得、検証、ビルド、公開、スモークテスト、障害切り分け。",
    ])
    add_callout(doc, "秘密情報の扱い", "本書にはパスワード、Firebaseの秘密鍵、API token、個人の2要素認証復旧コードは記載していません。引継ぎ時は各サービスの正式なメンバー招待・所有権移管を使い、秘密情報をチャットやリポジトリへ直接貼らないでください。", "info")


def add_exec_summary(doc):
    doc.add_heading("1. エグゼクティブサマリー", level=1)
    add_table(doc, ["観点", "現在地", "引継ぎ判断"], [
        ("プロジェクト", "Unity 6000.4.0f1 / WebGL / URP。5 Sceneが有効。", "同一Unity版とWebGL moduleを用意。"),
        ("体験", "4台が同じroomIdを共有。Pad1素材、Pad2作業台、Pad3 QR/能力、Pad4ホログラム。", "4台同時の実機確認が必要。"),
        ("状態", "RoomState schema v5。Firebase RESTまたはEditor内ローカル同期。", "Firebase所有権・Rules・データを別途移管。"),
        ("データ", "素材10、武器8、Rules 1、LaunchConfig 1、Pad4Calibration 1。", "定義Assetと.metaを必ず一緒に保全。"),
        ("検証", "基準検証0 error/1 warning。本番設定0 error。EditMode 134中132成功/2失敗。", "2失敗は旧Scene欠損Prefabと旧Step2期待値。"),
        ("Git", "追跡66、未追跡527（本書関連を除く）、LFS追跡実体0件。", "P0。現状のGitHub mainは復元元にならない。"),
        ("ビルド", "最後の確認済みWebGL buildは2026-07-30、13.24 MiB、成功。", "その後の変更を含む再ビルドが必要。"),
    ], [1550, 3900, 3910], font_size=8.0)

    doc.add_heading("引継ぎ完了の定義", level=2)
    add_bullets(doc, [
        "新アカウントでGitHub repositoryを閲覧・pushでき、Git LFSを含む全ソースがcloneできる。",
        "新アカウントがFirebase/Google Cloud projectとRealtime Database Rules/データを管理できる。",
        "Unity 6000.4.0f1で開き、基準検証・本番検証・テスト・WebGL buildを再実行できる。",
        "HTTPS公開URLで4つのscreenを同じroomに接続し、QR・転送・配置・合成・最終コードまで通る。",
        "旧アカウントを外しても、ソース・データ・ドメイン・端末・運用記録が残る。",
    ], numbered=True)


def add_git_critical(doc):
    doc.add_heading("2. 最重要：Git/GitHubだけでは現行プロジェクトを復元できない", level=1)
    add_callout(doc, "P0 / 即時対応", "2026-08-08〜09の監査時点で、git ls-filesは66件、未追跡は527件（本書用Tools/DocumentationとDocumentationを除く）でした。Assets配下の未追跡は523件、未追跡C#は78件、Git LFSの追跡実体は0件です。Craft-liveのScene、Script、定義、Prefab、3Dモデルの大部分がGitHubに存在しない状態です。", "critical")
    doc.add_heading("現在確認できるGit情報", level=2)
    add_table(doc, ["項目", "現在値", "意味"], [
        ("Remote", "https://github.com/goodmorning2424/CraftOrigin.git", "現所有者候補。新アカウントを管理者として招待し、必要ならrepository transfer。"),
        ("Branch", "main", "現行作業ツリーとmainの差が極端に大きい。"),
        ("Tracked", "66 files", "初期URP templateとPackages/ProjectSettings中心。"),
        ("Untracked", "527 files（引継ぎ生成物を除外）", "現行実装の中心。紛失リスク。"),
        ("Git LFS rules", ".gitattributesにobj/fbx/png等の指定あり", "ルールはあるが、git lfs ls-filesは0件。実体がまだ登録されていない。"),
        ("Ignored", "Library/Temp/Obj/Build/Builds/Logs/UserSettings等", "通常は正しい。復元元には含めず再生成する。"),
    ], [1500, 3300, 4560], font_size=8.0)
    doc.add_heading("安全な保全手順", level=2)
    add_bullets(doc, [
        ("1. 作業凍結：", "移管用snapshot作成中はUnity編集とファイル移動を止める。"),
        ("2. 外部バックアップ：", "Assets、Packages、ProjectSettings、.gitignore、.gitattributes、Tools、必要なREADMEを別媒体へコピーする。.metaは絶対に除外しない。"),
        ("3. Git LFS確認：", "LFS clientを導入し、.gitattributesの対象を含めて意図したファイルだけをstageする。"),
        ("4. 差分レビュー：", "Library/Logs/Builds/UserSettingsを誤って追加していないこと、秘密情報がないことを確認する。"),
        ("5. commit/push：", "引継ぎ用branchまたはmainへcommitし、LFS object upload完了まで確認する。"),
        ("6. 別PC復元試験：", "新規フォルダにclone + git lfs pullし、Unityで開いて検証・ビルドする。"),
        ("7. その後に所有権移管：", "clone検証が成功してからGitHub repositoryのowner/collaboratorを変更する。"),
    ], numbered=False)
    add_callout(doc, "本書ではcommit/pushを実行していません", "ユーザーの明示依頼がないため、作業ツリーの所有物を保護する目的でGitのstage/commit/pushやrepository transferは行っていません。", "warning")


def add_architecture(doc, diagrams):
    doc.add_heading("3. プロジェクト全体像", level=1)
    add_picture_with_alt(doc, diagrams["architecture"], 6.8, "Pad1〜Pad4がFirebase Realtime DatabaseのRoomStateを共有し、Bootstrapが各役割Sceneをロードする現行アーキテクチャ図")
    add_caption(doc, "図1　4つのPad Sceneと共通状態・共通ランタイム")
    doc.add_heading("起動URLと役割", level=2)
    add_table(doc, ["role", "screen/pad別名", "Scene", "主目的"], [
        ("MaterialPad", "items / materials / pad1", "Pad1_MaterialGallery", "素材閲覧、QR登録済み判定、候補選択、転送キュー、発射。"),
        ("Workbench", "craft / workbench / pad2", "Pad2_Workbench", "武器選択、6スロット、転送受信、液体、ハンマー合成、結果。"),
        ("Qr", "status / qr / pad3", "Pad3_StatusQr", "QR登録、メッセージ、攻撃/防御/回避の3管。"),
        ("Hologram", "hologram / pad4", "Pad4_Hologram", "完成武器、履歴、最終選択、CLコード。"),
    ], [1300, 2100, 2100, 3860], font_size=7.8)
    p = doc.add_paragraph("例：")
    p.add_run("https://<host>/?screen=items&room=room001").font.name = "Consolas"
    p.add_run(" / ")
    p.add_run("?screen=craft&room=room001").font.name = "Consolas"
    p.style = "Small Note"
    add_callout(doc, "ルーム分離", "roomパラメータが同じ4台だけが同じRoomStateを共有します。イベントごとにroom IDを分け、公開URLのキャッシュやブラウザ復元で別ルームへ入っていないか確認してください。", "info")

    doc.add_heading("Build Settings（有効順）", level=2)
    add_table(doc, ["Index", "Scene", "役割"], [
        (0, "CraftLiveBootstrap", "共通起動Scene"),
        (1, "Pad1_MaterialGallery", "素材画廊"),
        (2, "Pad2_Workbench", "作業台"),
        (3, "Pad3_StatusQr", "QR/能力"),
        (4, "Pad4_Hologram", "ホログラム"),
    ], [800, 3400, 5160], font_size=8.1)
    p = doc.add_paragraph("SampleSceneとCraft.unityは無効。Craft.unityは旧1画面構成で欠損Prefab参照があり、現行Build経路へ戻さない。PreStep2は移行前バックアップ。")
    p.style = "Small Note"


def add_flow_and_state(doc, diagrams):
    doc.add_heading("4. 体験フローと状態遷移", level=1)
    add_picture_with_alt(doc, diagrams["flow"], 6.8, "ルーム開始から素材登録、武器選択、転送、配置、合成、結果共有、最終決定へ進む体験フロー")
    add_caption(doc, "図2　標準体験フロー。赤字は現行Rulesによる重要な挙動")
    doc.add_heading("配置フローの状態", level=2)
    add_table(doc, ["値", "意味", "主な遷移"], [
        ("Idle", "待機・解除可能", "素材候補選択へ"),
        ("SelectingSlot", "配置先候補を選択", "ConfirmingSlotへ"),
        ("ConfirmingSlot", "選択を確認", "Pad1Loadingへ"),
        ("Pad1Loading", "発射装置へ装填", "Pad1Launchingへ"),
        ("Pad1Launching", "Pad1から飛翔", "Pad2Arrivingへ"),
        ("Pad2Arriving", "Pad2で到着演出", "PlacementCompleteへ"),
        ("PlacementComplete", "枠へ確定", "Idleへ戻り次素材"),
    ], [1650, 3100, 4610], font_size=8.1)
    doc.add_heading("合成・セッション状態", level=2)
    add_table(doc, ["区分", "Enum", "意味"], [
        ("Craft", "Editing", "武器/素材の編集、転送、配置。"),
        ("Craft", "Mixing", "液体/ハンマー合成中。"),
        ("Craft", "Complete", "結果が履歴へ追加済み。"),
        ("Session", "Playing", "300秒の通常体験。"),
        ("Session", "FinalSelection", "履歴から最終武器を選ぶ。"),
        ("Session", "Finished", "最終コード確定。"),
    ], [1300, 2000, 6060], font_size=8.1)
    doc.add_heading("Watchdogによる停止復旧", level=2)
    add_bullets(doc, [
        "Pad1側の転送状態が約6秒停滞すると次段階へ進める。",
        "Pad2Arrivingが約6秒停滞すると到着処理を補助する。",
        "配置完了表示は約3秒後に待機へ戻す。",
        "これは表示/入力事故に対する救済であり、Firebase切断時の競合マージを保証する機能ではない。",
    ])


def add_accounts(doc, diagrams):
    doc.add_heading("5. アカウント・外部サービス移管", level=1)
    add_picture_with_alt(doc, diagrams["ownership"], 6.8, "プロジェクトファイルで渡る情報と、GitHub、Firebase、Unity、Meshy、ホスティング等のアカウント側で移管する情報の境界")
    add_caption(doc, "図3　ファイルコピーだけでは完了しない引継ぎ境界")
    rows = [
        ("GitHub", "origin: goodmorning2424/CraftOrigin", "新アカウントをAdmin招待→完全push/LFS復元確認→必要ならrepository transfer→旧権限削除。", "P0"),
        ("Git LFS", ".gitattributesあり / LFS追跡実体0", "LFS object upload、quota/billing、新cloneで実体取得を確認。", "P0"),
        ("Firebase / Google Cloud", "craft-live-default-rtdb.firebaseio.com", "Project IAMのOwner/Editor、Realtime Database Rules、データexport、請求、ログ、リージョンを移管。", "P0"),
        ("Firebase認証", "クライアントはtoken/headerなしREST", "Rulesが認証必須なら現行クライアントは失敗。公開Rulesならroom URLを知る第三者のread/writeリスク。実Rulesを監査。", "P0"),
        ("Web hosting", "repositoryから提供者を特定できない", "公開先、deploy権限、Brotli header、WASM MIME、TLS、cache purge、rollbackを記録。", "P0"),
        ("Domain/DNS", "repositoryに所有情報なし", "registrar/DNS/CDNのOwner、2FA、更新課金、証明書を移管。", "P1"),
        ("Unity ID", "Unity Personalで検証 / Cloud services無効", "6000.4.0f1利用権、Hub login、組織/Cloud Projectを使用する場合のroleを確認。", "P1"),
        ("Meshy", "Assets/ai.meshy v0.2.2 / local bridge", "Meshyアカウント/API、利用規約・生成物権利、Bridge接続を別途引き継ぐ。鍵はproject内にない。", "P1"),
        ("iPad / MDM / Wi-Fi", "project外", "端末台帳、Safari設定、カメラ許可、Wi-Fi、充電、キオスク、時刻同期、予備機を移管。", "P1"),
        ("jsDelivr", "QR fallbackがqr-scanner@1.4.2を実行時import", "CDN障害/CSP/オフライン対策。必要ならself-hostへ変更。", "P2"),
    ]
    add_table(doc, ["対象", "現状", "移管作業", "優先"], rows, [1500, 2500, 4560, 800], font_size=7.25)
    doc.add_heading("Firebaseで必ず書き出す情報", level=2)
    add_bullets(doc, [
        "Firebase project ID、Google Cloud project number、Realtime Database instance、region。",
        "Realtime Database Rulesの全文と更新日時。公開read/writeか、認証条件があるか。",
        "運用中のroomデータを残すか削除するか、バックアップ/exportの保存先。",
        "所有者、Editor、Billing Admin、障害通知先。旧アカウント削除前に新アカウントでRules編集とデータ閲覧を試す。",
        "本番/ステージングinstanceが別にある場合はLaunchConfigとの対応表。",
    ])


def add_setup_build(doc):
    doc.add_heading("6. 開発環境・起動・ビルド", level=1)
    doc.add_heading("必要環境", level=2)
    add_table(doc, ["項目", "現行値", "注意"], [
        ("Unity Editor", "6000.4.0f1 (8cf496087c8f)", "同一版をUnity Hubへ追加。WebGL Build Supportを含める。"),
        ("Render Pipeline", "URP 17.4.0", "Mobile profile / Pad2 Glow volumeを使用。"),
        ("Active target", "WebGL", "iPad Safari向け3:4 portrait。"),
        ("Input", "Input System 1.19.0 / activeInputHandler=1", "旧Input Manager前提の追加コードに注意。"),
        ("Test Framework", "1.6.0", "EditMode 134件を確認。"),
        ("Network", "HTTPS + Firebase + QR CDN fallback", "カメラAPIはsecure contextが必要。"),
    ], [1800, 2800, 4760], font_size=8.0)
    doc.add_heading("初回復元手順", level=2)
    add_bullets(doc, [
        "GitHubからcloneし、git lfs pullを実行。別途snapshotの場合はAssets/Packages/ProjectSettingsと全.metaを同一階層へ配置。",
        "Unity Hubで6000.4.0f1を指定して開く。Libraryは転送せず再生成させる。",
        "Consoleのcompile errorが0であることを確認。Craft-live検証メニューまたはbatch validatorを実行。",
        "Build Settingsで5 Sceneと順序を確認し、Active Build TargetをWebGLへ。",
        "EditMode testsを実行し、既知2失敗以外の増加がないことを確認。既知失敗も保守方針に沿って解消する。",
        "WebGL buildを新規出力し、HTTPS hostで4つのURLを開いてスモークテスト。",
    ], numbered=True)
    doc.add_heading("Player / WebGL設定", level=2)
    add_table(doc, ["設定", "値", "理由/運用"], [
        ("Company / Product", "goodmorning2424 / Craft-live", "Player identity。"),
        ("Version", "0.1.0", "引継ぎ後はbuild識別のため更新推奨。"),
        ("Default canvas", "768 × 1024", "3:4 portrait。"),
        ("Run in Background", "ON", "Pad間同期を止めない。"),
        ("Color Space", "Linear", "URP/Glow表示。"),
        ("Template", "PROJECT:CraftLive", "Safe Area、DPR cap、online/offline表示。"),
        ("Compression", "Brotli", "Host側Content-Encoding設定が必要。"),
        ("Data caching", "ON", "更新時のcache invalidationを確認。"),
        ("Decompression fallback", "ON", "host設定不備の救済だが性能低下に注意。"),
        ("Threads", "OFF", "SharedArrayBuffer/COOP/COEP不要構成。"),
        ("Memory", "initial 32MB / max 2048MB / growth mode 2", "端末でmemory pressureを監視。"),
    ], [2200, 2900, 4260], font_size=7.8)
    doc.add_heading("品質・描画", level=2)
    add_bullets(doc, [
        "WebGLはMobile quality profile。render scale 0.8、MSAA 1、depth/opaque texture OFF。",
        "Main light shadow ON（1024）、additional lightはper-vertexでshadowなし、shadow distance 50。",
        "SRP Batcher ON、Adaptive Performance ON。",
        "Pad2 Glow VolumeはBloom threshold 0.9、intensity 0.28、scatter 0.55。旧文書の「post processなし」は現行と不一致。",
    ])


def add_scenes(doc, scenes):
    doc.add_heading("7. Scene構成と現行Inspector要点", level=1)
    summary_rows = []
    for scene in scenes:
        path = scene["path"]
        kind = "有効" if "Assets/Scenes/CraftLive/" in path and "PreStep" not in path else "無効/履歴"
        summary_rows.append((Path(path).stem, kind, scene["game_object_count"], scene["root_count"], len(scene["source_prefabs"]), path))
    add_table(doc, ["Scene", "区分", "GO", "Root", "Prefab種", "Path"], summary_rows, [1600, 1200, 600, 600, 700, 4660], font_size=7.2)

    doc.add_heading("CraftLiveBootstrap", level=2)
    add_bullets(doc, [
        "共通rootにSession、RoomTransport、Bootstrap、WebPresentation、Timer、RuntimeDiagnostics。Camera + PhysicsRaycasterを1つずつ持つ。",
        "BootstrapがLaunchConfigを読み、URLのscreen/pad/roomを解決。Pad role SceneをAdditive Loadする。",
        "WebPresentation：target 30fps、vSync 0、sleep禁止、Safe Area、3:4 viewport。Diagnostics stale判定15秒。",
        "Transport Scene値のuseFirebase=falseはLaunchConfig/Bootstrapで実行時上書きされる。Editorはローカル、WebGLはFirebase。",
    ])

    doc.add_heading("Pad1_MaterialGallery", level=2)
    add_table(doc, ["機能", "現行値/参照", "引継ぎ注意"], [
        ("Gallery", "10素材、showLocked=false、2列、column 2.75、painting 2.15×1.72、drag 0.012、wheel 0.8", "Catalog変更後は額Prefab/spacing/壁範囲を実機確認。"),
        ("Preview", "target 0.9、rotation (10,-20)、spin 60°/s、reveal 0.28、group scale 0.46", "多くのworldPrefabがCube placeholder。"),
        ("Transfer", "queue 3 columns、required pull 110、launchAllByDefault=true", "同時発射時の順序と予約枠を確認。"),
        ("Durations", "load .28 / arrival .4 / impact .12 / groove .24 / launch .55", "低速端末でWatchdogとの競合を監視。"),
        ("Camera shift", "3 sec / yaw 71°", "Scene camera変更時に再調整。"),
        ("Play-test bypass", "allowTransferWithoutPlacementForPlayTest=true", "P1。本番で配置確認を迂回できるためfalse化を判断。"),
    ], [1700, 4150, 3510], font_size=7.6)

    doc.add_heading("Pad2_Workbench", level=2)
    add_table(doc, ["機能", "現行値/参照", "引継ぎ注意"], [
        ("Weapon carousel", "swipe 70、spacing 2.65、neighbor .62、selected 1.1、center 1.4", "全武器が同一icon。hologramPrefab未設定。"),
        ("Placement", "6論理slot、fallback controls、reference layout、diameter .86", "Slot enum番号はserialized contract。並び替え禁止。"),
        ("Receiver", "delay .35、duration .75、arc 1.2、surface .45、hold 3.2", "publishStatsAfterArrival=false。配置確定時に反映。"),
        ("Liquid", "prefabなしfallback、6 drops、flow 2.4、radius .11、pool .38", "演出差替えはPrefab参照と性能を確認。"),
        ("Hammer", "fallback hammer、rail 2.6、6 passes", "stroke 120。入力閾値と端末タッチを確認。"),
        ("Result", "fallback visual", "attribute/skill display ID空欄では表示が欠ける。"),
    ], [1700, 4150, 3510], font_size=7.6)

    doc.add_heading("Pad3_StatusQr", level=2)
    add_bullets(doc, [
        "CraftLiveQrScannerのScene timeoutは8秒。WebGL側はカメラオーバーレイ、結果をSendMessageで返す。",
        "3本のStatusTube：fullHeight 3.2、bottom -1.6、width .42、animation .55。Controllerが攻撃=赤、防御=青、回避=緑を適用。",
        "登録は消費型在庫ではなく永久登録。QR必須素材はregisteredMaterialIdsに入ると以後使用可能。",
        "HTTPS、Safariカメラ許可、背面カメラ選択、CSP/CDN到達性を実機確認。",
    ])

    doc.add_heading("Pad4_Hologram", level=2)
    add_bullets(doc, [
        "Hologram rotate=true、30°/s、属性色反映ON、emission 2。表示Prefab未設定時はworkbenchPrefabへfallback。",
        "Calibration：画面180×240mm、角度45°、視距離80mm、位置/回転0、scale 1。端末/ミラー筐体変更時は再計測。",
        "履歴最大12件から最終選択し、XXXXXX形式の6文字コードを表示する。",
    ])

    doc.add_heading("無効Sceneの扱い", level=2)
    add_table(doc, ["Scene", "扱い"], [
        ("Craft.unity", "旧1 Scene構成。欠損Prefab GUID 01b025…を含み、最新テスト失敗要因。Buildへ再登録しない。修復するかアーカイブ方針を決める。"),
        ("SampleScene.unity", "URP template由来。現行体験では未使用。"),
        ("PreStep2_Craft", "Step2移行前バックアップ。現行仕様の参照元にしない。"),
    ], [2300, 7060], font_size=8.0)


def add_state_model(doc):
    doc.add_heading("8. RoomState schema v5と永続データ", level=1)
    add_table(doc, ["領域", "主要フィールド", "役割"], [
        ("Versioning", "schemaVersion / revision / updatedAt", "互換性、競合比較、stale判定。"),
        ("Registration", "registeredMaterialIds / selectedMaterialId", "QR素材の永久登録とPad1選択。"),
        ("Legacy", "inventory / qrUnlocked", "V1/V2移行互換。Normalize後は登録へ寄せる。"),
        ("Placement", "status / materialId / targetSlot / serial", "転送・配置の単一進行状態。"),
        ("Transfer queue", "entries / nextSerial / batchRemaining", "Pad1からPad2へ複数素材を順序付け。"),
        ("Registration event", "registeredMaterialId / registeredAt / serial", "Pad3登録通知。"),
        ("Slots", "attribute / skill / top / right / left / bottom", "6枠のmaterialId。二重予約を防止。"),
        ("Displayed stats", "displayedStats / statusSerial", "Pad3表示用。武器確定・解除・結果で更新。"),
        ("Craft", "status / hammerPasses / power / startedAt", "編集→合成→完了。"),
        ("Weapon", "selectedWeaponId / weaponConfirmed", "Pad2武器状態。"),
        ("Result", "current result + completedResults", "完成武器と最大12件の履歴。"),
        ("Session", "phase / startedAt / expiresAt", "Playing / FinalSelection / Finished。"),
        ("Final", "selectedResultSerial / code / message", "最終武器とCLコード。"),
    ], [1650, 3500, 4210], font_size=7.7)
    doc.add_heading("serialized contractとして固定すべきEnum値", level=2)
    add_table(doc, ["Enum", "値"], [
        ("CraftLiveRole", "Auto=0, MaterialPad=1, Workbench=2, Qr=3, Hologram=4"),
        ("MaterialCategory", "Attribute=0, Skill=1, Upgrade=2"),
        ("SlotId", "Attribute=0, Skill=1, Top=2, Right=3, Left=4, Bottom=5"),
        ("StatType", "None=0, Attack=1, Defense=2, Evasion=3, ElementBoost=4（legacy/reserved）"),
        ("ElementType", "None=0, Fire=1, Freeze=2, Lightning=3"),
        ("SkillType", "None=0, Luck=1, DoubleStrike=2, AutoHeal=3, LifeOrb=4"),
        ("WeaponType", "Sword=0, Thrust=1, Staff=2"),
    ], [2100, 7260], font_size=8.0)
    add_callout(doc, "Enumの並び替え禁止", "Unity YAMLとRoomState JSONでは整数値として保存される箇所があります。既存値の順番変更や途中挿入は旧データ・Scene参照を壊します。追加は末尾、変更時は明示的migrationを実装してください。", "warning")
    doc.add_heading("Normalize / migration", level=2)
    add_bullets(doc, [
        "schema v1/v2のinventory・qrUnlockedを読み取り、登録済み素材へ統合する。",
        "現在は登録素材を消費しない。互換Countは登録済みなら1を返す。",
        "null list/object、無効slot、重複ID、履歴上限を正規化する。",
        "将来schemaを上げる場合は旧JSON fixtureを残し、全Padを同時deployできない時間帯の前方/後方互換を設計する。",
    ])


def add_transport(doc):
    doc.add_heading("9. 同期・Firebase・競合・オフライン", level=1)
    add_table(doc, ["項目", "現行実装"], [
        ("Editor/local", "static dictionaryで同一Editor process内のroomを共有。"),
        ("WebGL/remote", "https://craft-live-default-rtdb.firebaseio.com/rooms/{roomId}.json をREST polling。"),
        ("Poll / timeout", "0.5秒 / 10秒。"),
        ("Write", "GETでETag取得、PUTにIf-Matchを付与。"),
        ("新旧比較", "revisionを優先、同値ならupdatedAt。"),
        ("Conflict", "412等で再読込。field-level mergeはしない。"),
        ("Retry", "0.75秒から8秒まで指数backoff。"),
        ("Connection state", "失敗が続くとDegraded、3回でOffline。"),
        ("Local cache", "PlayerPrefsにpending/confirmed RoomState。pending cache有効。"),
    ], [2200, 7160], font_size=8.0)
    add_callout(doc, "競合の限界", "ETagは同時上書きを検出しますが、再読込後に個別フィールドを自動マージしません。長時間オフラインのPadが復帰して古い意図を再送すると、別Padの変更を失う可能性があります。イベント運用では同一roomへの重複操作を避け、復帰時にrevisionと表示を確認してください。", "warning")
    doc.add_heading("Firebase Rulesのセキュリティ判断", level=2)
    add_bullets(doc, [
        "現行クライアントはAuth token/API key/Authorization headerを送っていない。",
        "Rulesがpublic read/writeなら動作するが、URLとroomIdを知る第三者がRoomStateを読書きできる。",
        "Rulesが認証を要求するなら現行クライアントはHTTPエラーとなる。",
        "本番では期限付きroom、App Check/Auth導入、server relay、IP/運用制限等を選択し、クライアント改修と同時にRulesを更新する。",
        "少なくとも引継ぎ時にRulesをexportし、実値を本書の別管理票へ記録する。",
    ])


def add_pads_detail(doc):
    doc.add_heading("10. Pad別の詳細責務と操作", level=1)
    pad_sections = [
        ("Pad1：素材画廊・転送", [
            "Catalogの10素材を額装表示。QR必須かつ未登録の素材は使用不可表示にする。showLocked=falseのため未登録作品を隠す構成。",
            "作品選択で3D previewと説明を表示。素材定義のworldPrefabが未設定/placeholderの場合はfallback visual。",
            "転送はtarget slot予約→確認→queue→バネ引き→load/launch→Pad2 arrivalの順。予約済みslotを別素材が使えない。",
            "launchAllByDefault=trueのためキュー内をまとめて送れる。RoomState serialで順序を保持。",
            "解除はPlacement Idle時のみ。解除直後にdisplayedStatsを再計算・publish。",
        ]),
        ("Pad2：武器・配置・合成", [
            "8武器カルーセル。確定時にweaponConfirmed=true、武器基礎能力をdisplayedStatsへ反映。",
            "スロットはAttribute、Skill、Top、Right、Left、Bottom。Attribute/Skillは効果、4基礎枠はstatModifiersに使用。",
            "Pad1キューを監視し、arrival arc→surface→groove→slot確定を表示。",
            "現行RulesではrequireAttribute/requireSkill/requireAllFourBaseSlotsが全てfalse。武器確定だけで合成可能。",
            "ハンマー6打。progress×100×lerp(0.75,1,quality)の最大値をpowerとし、6打で自動完了。",
            "結果をcompletedResultsへ追加し、新しい武器へ進むとslotsをreset。ただし登録素材と履歴は保持。",
        ]),
        ("Pad3：QR・能力", [
            "WebGL JavaScriptへQR開始を要求。BarcodeDetectorが利用できればnative、なければqr-scanner 1.4.2をCDN import。",
            "QR文字列を素材IDとして解決し、登録済みIDへ追加。登録は永久で消費しない。",
            "攻撃/防御/回避を0〜100で3本の液面に表示。武器確定、slot解除、完成結果で更新。",
            "カメラ許可拒否、HTTPSでない、CSP/CDN blocked、別camera選択などを運用時に切り分ける。",
        ]),
        ("Pad4：ホログラム・最終決定", [
            "current resultまたは選択履歴の武器Prefabを表示。hologramPrefabが空ならworkbenchPrefabを使用。",
            "30°/s回転、属性色とemissionを適用。Calibrationで物理表示面へ合わせる。",
            "セッション終了後、最大12件の完成履歴から1件を選択。",
            "6文字で武器、属性、スキル、攻撃・防御・回避素材数を直接表す。ルーム、時刻、serial、hashは含まない。",
        ]),
    ]
    for title, bullets in pad_sections:
        doc.add_heading(title, level=2)
        add_bullets(doc, bullets)


MATERIALS = [
    ("ore_attack", "攻撃", "Upgrade", "不要", "+5/0/0", "Attack prefab / Atttack.png", "現行基礎slot用。filename/iconの綴りに注意。"),
    ("ore_defence", "防御", "Upgrade", "不要", "0/+5/0", "Defence prefab/icon", "現行基礎slot用。"),
    ("ore_evasion", "回避率", "Upgrade", "不要", "0/0/+5", "Kaihi prefab/icon", "現行基礎slot用。"),
    ("attribute_fire", "炎の結晶", "Attribute", "必要", "Fire / chance 0", "Cube / Fire.png", "attributeId=fire、displayName=炎。chance 0でvalidator warning対象。"),
    ("attribute_freeze", "氷", "Attribute", "必要", "Freeze / chance 0", "Cube / Snow.png", "attributeId/displayName空。旧defense+30はカテゴリ上計算対象外。warning対象。"),
    ("attribute_lighting", "雷", "Attribute", "必要", "Lightning / 48.9%", "Cube / Sunder.png", "lighting綴り、attributeId/displayName空。"),
    ("skill_double_strike", "二回攻撃", "Skill", "必要", "DoubleStrike / 0%", "Cube / Wind.png", "skillId/displayNameあり。type Noneでもresolver推定。"),
    ("skill_auto_heal", "自動回復", "Skill", "必要", "AutoHeal / 0%", "Cube / Heal.png", "skillId/displayName空。旧attack+30はカテゴリ上計算対象外。"),
    ("skill_life_orb", "命の珠", "Skill", "必要", "LifeOrb / 50.2%", "Cube / Life.png", "skillId/displayName空。"),
    ("skill_luck", "幸運", "Skill", "必要", "Luck / 31.8%", "Cube / root icon", "skillId/displayName空。"),
]

WEAPONS = [
    ("weapon_bigsword", "大剣", "Sword", "50/50/10", "Taiken"),
    ("weapon_fude_staff", "ふで", "Staff", "80/40/30", "Fude"),
    ("weapon_katate", "片手剣", "Sword", "40/20/50", "Katate"),
    ("weapon_kaziki", "カジキマグロ", "Thrust", "100/10/10", "Kaziki"),
    ("weapon_kobushi", "こぶし", "Thrust", "90/90/90", "Kobushi"),
    ("weapon_pikopiko", "ピコピコハンマー", "Sword", "10/10/100", "Pikopiko"),
    ("weapon_tue", "杖", "Staff", "60/60/20", "Tue"),
    ("weapon_rapier", "やり", "Thrust", "50/10/50", "Yari"),
]


def add_catalog(doc):
    doc.add_heading("11. 現行データCatalog（素材10 / 武器8）", level=1)
    doc.add_heading("素材Definition", level=2)
    add_table(doc, ["ID", "表示名", "Category", "QR", "能力/効果", "Visual", "注意"], MATERIALS, [1450, 1050, 900, 550, 1450, 1700, 2260], font_size=6.85)
    add_callout(doc, "素材データの未完成点", "2026-08-08 validatorは2 definitionsをwarning（fire/freezeのchance 0等）と判定しています。さらにfreeze/lightingおよび複数skillでhuman-readable ID/displayNameが空です。結果名・最終code入力・表示が欠ける可能性があるため、仕様確定後に全10素材を一括整備してください。", "warning")
    doc.add_heading("武器Definition", level=2)
    rows = [(i, n, t, s, p, "hologramPrefab空 / 共通Atttack.png / scale 1") for i, n, t, s, p in WEAPONS]
    add_table(doc, ["ID", "表示名", "Type", "攻/防/避", "Workbench Prefab", "共通注意"], rows, [1900, 1250, 900, 1100, 1700, 2510], font_size=7.2)
    add_bullets(doc, [
        "weapon_rapierの表示/Prefabは「やり/Yari」で、IDと意味が不一致。外部データ連携前にID変更か互換aliasを決める。",
        "全武器のhologramPrefabが空のためPad4はworkbenchPrefabへfallback。専用軽量Prefabを作ると表示と性能を分離できる。",
        "全武器が同じAtttack.pngを参照。識別性とアクセシビリティ改善のため個別iconが望ましい。",
    ])


def add_rules_calc(doc):
    doc.add_heading("12. ルール・計算・コード生成", level=1)
    add_table(doc, ["Rule", "現行値", "影響"], [
        ("sessionDurationSeconds", "300", "5分でFinalSelectionへ。"),
        ("requireAttribute", "false", "属性なしでも合成可。"),
        ("requireSkill", "false", "スキルなしでも合成可。"),
        ("requireAllFourBaseSlots", "false", "基礎4枠が空でも合成可。"),
        ("mixingDuration", "5", "旧/補助値。現行主制御はhammer 6 passes。"),
        ("powerPerRadian", "7.5", "旧MixInput用。"),
        ("hammerPasses / stroke", "6 / 120", "6打で自動確定。"),
        ("history", "12", "completedResults最大件数。"),
        ("rank thresholds", "31 / 61 / 91", "Success / Great / Superの境界。"),
        ("rank bonuses", "+5 / +10 / +15", "攻防避すべてへ加算。"),
        ("maxStat", "100", "最終値をclamp。"),
        ("codePrefix", "CL", "旧形式との互換用。現行コードには使用しない。"),
    ], [2700, 2200, 4460], font_size=8.0)
    doc.add_heading("ステータス計算", level=2)
    add_bullets(doc, [
        "開始値＝確定武器のbase attack/defense/evasion。",
        "Top/Right/Left/Bottomの4基礎slotにあるUpgrade素材のstatModifiersを加算。Attribute/Skillカテゴリのlegacy statModifiersは計算に使わない。",
        "hammer powerが31未満は通常成功でbonus 0、31以上+5、61以上+10、91以上+15。",
        "ランクbonusを攻撃・防御・回避へ同値加算し、各値を0〜100へclamp。",
        "AttributeとSkillの効果はStatsとは別にresultへ保存し、Pad4表示・code生成へ使う。",
    ], numbered=True)
    doc.add_heading("ハンマーPower", level=2)
    p = doc.add_paragraph("各打撃の候補Power = progress × 100 × lerp(0.75, 1.0, quality)。セッション中のPowerは候補の最大値。6打完了でランクと結果を確定する。")
    p.paragraph_format.left_indent = Inches(0.25)
    add_callout(doc, "仕様確認ポイント", "現行Rulesでは素材0件でも合成できます。旧文書の「属性・スキル・4基礎枠が必須」を採用したい場合は、3つのrequireフラグを変更し、UIの案内/テスト/運用を同時更新してください。", "warning")


def add_web_qr(doc):
    doc.add_heading("13. WebGLテンプレート・QR・公開", level=1)
    add_table(doc, ["ファイル/機能", "内容", "外部条件"], [
        ("WebGLTemplates/CraftLive/index.html", "Unity loader、viewport、Safe Area、DPR cap 1.5、online/offline表示。", "HTTPS host、正しい相対path。"),
        ("style.css", "full viewport / safe-area-inset / portrait表示。", "iOS Safariのviewport変化を確認。"),
        ("CraftLiveWebGL.jslib", "URL query読取、QR overlay、BarcodeDetector、SendMessage callback。", "カメラ許可、secure context。"),
        ("qr-scanner fallback", "jsDelivrからv1.4.2をdynamic import。", "Internet/CSP/CDN。"),
        ("Firebase REST", "rooms/{roomId}.jsonへGET/PUT。", "Rules、CORS、HTTPS、接続。"),
    ], [2500, 4100, 2760], font_size=7.7)
    doc.add_heading("公開サーバー設定チェック", level=2)
    add_bullets(doc, [
        "HTTPSを必須化。HTTP→HTTPS redirectを確認。QR cameraは非secure contextでは使えない。",
        "Brotli圧縮ファイルへContent-Encoding: br、.wasmへapplication/wasm、.jsへJavaScript MIMEを設定。",
        "新build公開時はHTML/loaderとhash付きdata/wasmのcache policyを分け、古いService Worker/CDN cacheをpurge。",
        "CSPを設定する場合、Firebase endpoint、jsdelivr dynamic import、camera/mediaを許可するか、QR scriptをself-host。",
        "4 URLを別tab/別iPadで同じroomに接続。roomをURL encodeし、推測困難なevent IDを使う。",
        "最終build artifact、deploy commit、公開日時、rollback先を運用記録へ残す。",
    ], numbered=True)
    add_callout(doc, "最新buildの鮮度", "確認できるBuilds/CraftLiveWebGLは2026-07-30作成で13.24 MiB、成功（353 warnings）です。その後のScene/データ/コード変更を含まない可能性が高いため、引継ぎ用の正本として再ビルド・再スモークテストしてください。Buildsは.gitignore対象です。", "warning")


def add_project_structure(doc, manifest):
    doc.add_heading("14. フォルダ構成・パッケージ・容量", level=1)
    top_rows = []
    for path in sorted((ROOT / "Assets").iterdir()):
        if path.name.endswith(".meta"):
            continue
        files = [p for p in path.rglob("*") if p.is_file()] if path.is_dir() else [path]
        non_meta = [p for p in files if p.suffix.lower() != ".meta"]
        total = sum(p.stat().st_size for p in non_meta)
        top_rows.append((path.name, len(non_meta), fmt_size(total), {
            "Scripts": "Craft-live code/docs",
            "Scenes": "enabled + legacy scenes",
            "CraftLiveData": "catalog/definition/rules/config/calibration + visual data",
            "Buki": "weapon meshes/prefabs",
            "Pad1": "gallery/launcher models",
            "Materials": "Unity materials",
            "Editor": "validators/upgraders/tests",
            "Plugins": "WebGL QR bridge",
            "WebGLTemplates": "custom player template",
            "ai.meshy": "Meshy Unity bridge",
        }.get(path.name, "supporting assets")))
    add_table(doc, ["Assets直下", "非meta件数", "容量", "役割"], top_rows, [2000, 1200, 1300, 4860], font_size=7.4)
    doc.add_heading("直接依存パッケージ", level=2)
    package_json = json.loads((ROOT / "Packages" / "manifest.json").read_text(encoding="utf-8"))
    pkg_rows = [(name, version, "Runtime/Editor" if name in ("com.unity.inputsystem", "com.unity.render-pipelines.universal", "com.unity.ugui") else "Editor/support") for name, version in package_json.get("dependencies", {}).items()]
    add_table(doc, ["Package", "Version", "主用途"], pkg_rows, [4700, 2200, 2460], font_size=7.2)
    add_bullets(doc, [
        "Unity AI Assistant 2.16.0-pre.1とAI Inference 2.6.1は開発支援/推論用依存。アカウント権限やpreview package更新に注意。",
        "Meshy pluginはPackagesではなくAssets/ai.meshyに同梱（v0.2.2、GPL-3.0）。runtime体験の必須依存ではない。",
        "Library/PackageCacheは復元時に再取得される。Packages/manifest.jsonとpackages-lock.jsonを正本として保全。",
    ])
    counts = Counter(ext for _, ext, _, _ in manifest if _.startswith("Assets/"))
    doc.add_heading("保全対象と再生成対象", level=2)
    add_table(doc, ["区分", "対象", "方針"], [
        ("必須正本", "Assets（全.meta含む）、Packages、ProjectSettings", "Git/LFSまたはsnapshotで保全。"),
        ("推奨", ".gitignore、.gitattributes、Tools、文書、deploy設定", "復旧/運用の再現に必要。"),
        ("再生成", "Library、Temp、Obj、Logs", "容量が大きく端末依存。原則転送しない。"),
        ("成果物", "Builds/CraftLiveWebGL", "ソース正本ではないが、rollback artifactとして別保管。"),
        ("要判断", "LocalBackups、GeneratedAssets、UserSettings", "必要な原本のみ選別し、個人設定/重複を正本化しない。"),
    ], [1300, 3900, 4160], font_size=8.0)


def add_code_inventory(doc, scripts):
    doc.add_heading("15. コード構成と責務", level=1)
    layers = Counter()
    rows = []
    for s in scripts:
        parts = s["path"].split("/")
        if "Editor" in parts:
            layer = "Editor/Test"
        else:
            idx = parts.index("CraftLive") if "CraftLive" in parts else -1
            layer = parts[idx + 1] if idx >= 0 and idx + 1 < len(parts) else "Other"
        layers[layer] += 1
        rows.append((layer, s["name"], ", ".join(s["decls"]) or "-", s["lines"], infer_script_description(s["name"], s["path"])))
    add_table(doc, ["Layer", "File", "主要宣言", "行", "責務"], rows, [1050, 1900, 2200, 600, 3610], font_size=6.65)
    add_callout(doc, "巨大Viewクラス", "Pad1MaterialPreviewやPad1TransferControllerは2,000行を超える大規模クラスです。UI生成・状態反映・入力・演出が混在しているため、変更時はScene参照とfallback生成の両経路をテストしてください。", "info")
    doc.add_heading("レイヤー依存の基本", level=2)
    add_bullets(doc, [
        "CoreはUnity Sceneに依存しない計算・状態契約。最優先で単体テストを維持する。",
        "DataはScriptableObjectによる調整値とAsset参照。IDの一意性・空欄・category整合をvalidatorで守る。",
        "RuntimeはSession/Transport/Bootstrap/Timer。全Pad共通の状態遷移をここへ集約する。",
        "UI/Viewは入力と表示。SessionのRoomStateを直接壊さず公開API経由で更新する。",
        "Editor/TestはScene生成・移行履歴と検証。旧Step前提のテストは現行仕様へ更新またはlegacy suiteへ分離する。",
    ])


def add_quality_risks(doc):
    doc.add_heading("16. 検証結果・既知課題・優先順位", level=1)
    add_table(doc, ["検証", "実行日時", "結果", "解釈"], [
        ("Current Project Validation", "2026-08-08 23:43 JST", "Errors 0 / Warnings 1", "Scene参照・Catalog・schema・build sceneは合格。2 definition不完全。"),
        ("Production Readiness", "2026-08-08 23:44 JST", "Errors 0", "WebGL/Template/QR/Firebase URL/cache/retry構成は合格。外部host/Firebase Rulesは対象外。"),
        ("EditMode Tests", "2026-08-08 23:38–23:39 JST", "134 total / 132 pass / 2 fail / 0 skip", "旧Craft.unity欠損Prefabと旧Step2 transport.enabled期待値。"),
        ("Historical WebGL Build", "2026-07-30", "Succeeded / 13.24 MiB / 353 warnings", "最新変更前。引継ぎ版を再build。"),
    ], [2100, 1900, 2300, 3060], font_size=7.7)
    risks = [
        ("P0", "Git未追跡527 / LFS 0", "GitHubだけでは復元不能。", "完全snapshot→意図したcommit/LFS push→別PC clone検証。"),
        ("P0", "Firebase/hosting所有権不明", "ソースがあっても同期・公開を継続できない。", "IAM/Rules/data/DNS/hostを正式移管。"),
        ("P0", "Firebase REST無認証", "public Rulesなら改ざん、privateなら動作不能。", "実Rules監査と脅威モデル確定。"),
        ("P1", "fire/freeze gameplay values不完全", "validator warning、効果が発動しない。", "chance/effect仕様とテストを確定。"),
        ("P1", "複数素材のID/displayName空", "結果表示/コード入力が欠ける。", "10素材の定義を一括補完。"),
        ("P1", "Pad1 play-test bypass=true", "本番制約を迂回。", "本番false化、テスト専用configへ分離。"),
        ("P1", "全武器hologramPrefab空/同一icon", "Pad4性能・識別性・見栄え。", "専用Prefab/iconとLODを作る。"),
        ("P1", "最新tests 2 fail", "回帰品質の基準が曖昧。", "legacy Scene修復/除外、Step2期待値を現行へ更新。"),
        ("P1", "WebGL buildが7/30", "現行コードの公開物ではない。", "新build、実機4台E2E、artifact保存。"),
        ("P2", "lighting/rapierのID不一致", "外部連携・保守で混乱。", "互換migrationまたはalias方針。"),
        ("P2", "QR CDN fallback", "CDN/CSP/オフライン障害。", "self-hostまたはnative限定判断。"),
        ("P2", "競合はfield mergeなし", "長時間offline復帰で更新損失。", "操作権/command log/transaction設計を検討。"),
    ]
    add_table(doc, ["優先", "課題", "影響", "推奨"], risks, [700, 2250, 2600, 3810], font_size=7.2, header_fill=PALE_RED)
    doc.add_heading("2件の失敗詳細", level=2)
    add_table(doc, ["Test", "原因", "現行への影響", "保守案"], [
        ("CraftScene_HasSixUniqueWorkbenchAnchors", "無効Craft.unityがmissing prefab GUID 01b025a9…を参照し、Scene open時にError log。", "Build経路では無効。旧Sceneを開く/再利用すると問題。", "旧Sceneを修復して履歴保存、またはlegacy test対象外を明示。"),
        ("BootstrapScene_HasRequiredRuntimeReferences", "Step2 testはtransport.enabled=falseを期待、現行Sceneはtrue。", "現行validator/本番検証はTransport参照を合格。", "現行起動仕様に合わせて期待値更新。Bootstrap上書きとの責務もテスト。"),
    ], [2500, 3000, 2000, 1860], font_size=7.25)


def add_handoff_sop(doc):
    doc.add_heading("17. 引継ぎ実施手順（SOP）", level=1)
    phases = [
        ("A. 保全", [
            "作業停止時刻を宣言し、現PCの完全snapshotを作る。",
            "Assets/Packages/ProjectSettingsの件数とhash manifestを保存。.meta欠落がないか確認。",
            "Gitへ意図的にstage/commitし、binaryはLFS object uploadを完了。",
            "Builds/最新artifactとFirebase Rules/exportはソースとは別の安全な保管先へ。",
        ]),
        ("B. アカウント移管", [
            "GitHub新所有者/Adminを追加。branch protection、Actions secret、LFS billingも確認。",
            "Firebase/Google Cloud IAM、Billing、Rules、Database、通知先を移管。",
            "Hosting/DNS/TLS、Unity org、Meshy、端末/MDM/Wi-Fiを移管。",
            "新アカウントで操作確認後、旧アカウントの権限を段階的に削除。",
        ]),
        ("C. 復元検証", [
            "別ディレクトリ/別PCでclone + LFS pull。既存Libraryを流用しない。",
            "Unity 6000.4.0f1でopenし、compile 0 error。",
            "CurrentValidation 0 error、ProductionReadiness 0 error、EditMode testを実行。",
            "新WebGL buildをstagingへdeployし、4台E2Eを実施。",
        ]),
        ("D. 切替", [
            "本番deploy、cache purge、4 URLのroom接続、QR camera、Firebase書込みを確認。",
            "旧deploy/旧databaseのrollback期限を決め、監視と問い合わせ窓口を更新。",
            "引継ぎ完了サイン：新owner、旧owner、技術担当、運用担当の4者。",
        ]),
    ]
    for title, items in phases:
        doc.add_heading(title, level=2)
        add_bullets(doc, items, numbered=True)

    doc.add_heading("4台E2Eスモークテスト", level=2)
    smoke = [
        (1, "起動", "4台をitems/craft/status/hologram、同一roomで開く。", "各role Scene、online、同room。"),
        (2, "QR", "Pad3でQR必須素材を登録。", "Pad1で使用可能になり、再読取しても消費されない。"),
        (3, "武器", "Pad2で任意の武器を選択・確定。", "Pad3能力が基礎値へ更新。"),
        (4, "素材", "Pad1で素材を選びslot指定、バネを引き発射。", "Pad2へ到着し同slotに確定。"),
        (5, "解除", "Idleでslotを解除。", "二重予約なし、Pad3能力が即時再計算。"),
        (6, "合成", "必要な構成で液体→ハンマー6打。", "Power/Rank/Stats/resultが一致。"),
        (7, "履歴", "2本以上作成。", "履歴保持、次武器でslot reset、登録は保持。"),
        (8, "時間", "300秒経過またはfinal selectionへ。", "Pad4で履歴選択可能。"),
        (9, "最終", "1件選び確定。", "6文字コード表示、全PadでFinished同期。"),
        (10, "障害", "1台をoffline→復帰。", "Degraded/Offline表示、復帰後revisionと表示が収束。"),
    ]
    add_table(doc, ["#", "区分", "操作", "合格条件"], smoke, [500, 1200, 3900, 3760], font_size=7.4)


def add_operations(doc):
    doc.add_heading("18. 運用・障害切り分け", level=1)
    add_table(doc, ["症状", "最初に確認", "次の確認/対応"], [
        ("Pad間で同期しない", "4台のroom/online/Firebase URL", "HTTP status、Rules、revision、device clock、cache。"),
        ("QR画面が出ない", "HTTPS、camera permission", "BarcodeDetector対応、CSP、jsDelivr、Safari設定。"),
        ("QRは読めるが登録されない", "QR文字列とmaterial ID", "Catalog、registeredMaterialIds、Firebase write error。"),
        ("素材が飛ばない", "placement status/queue/slot予約", "Pad1 bypass、spring threshold、Watchdog、Pad2 receiver。"),
        ("Pad2に届かない", "Pad2 role Sceneとsame room", "transfer serial、arrival status、offline cache。"),
        ("合成できない", "weaponConfirmed / Rules required flags", "slot ID/category、hammer input、CraftStatus。"),
        ("能力が想定外", "weapon base、4 base slots、rank bonus", "Attribute/Skillのlegacy modifiersを誤解していないか。"),
        ("Pad4がCube/武器違い", "hologramPrefab/workbenchPrefab", "ID、Catalog参照、fallback visual、Calibration。"),
        ("更新後に古い画面", "browser/CDN cache", "Build hash、HTML cache-control、data caching、purge。"),
        ("Sceneを開くとMissing Prefab", "対象が旧Craft.unityか", "現行5 Sceneならvalidator。旧Sceneはrepair/archive。"),
    ], [2200, 3200, 3960], font_size=7.6)
    doc.add_heading("ログ/証跡", level=2)
    add_bullets(doc, [
        "Unity Editor ConsoleとLogs/*.log。",
        "Library/CraftLiveReports/CurrentValidation_latest.md / ProductionReadiness_latest.md。",
        "EditMode test XML（total/passed/failed、失敗message/stack）。",
        "Firebase Realtime Database usage/audit/Rules history（利用可能な範囲）。",
        "Hosting access/error log、CDN purge/deploy履歴。",
        "イベントごとのroom ID、開始/終了時刻、担当者、端末番号。",
    ])


def add_existing_docs(doc):
    doc.add_heading("付録A. 既存文書の使い分け", level=1)
    rows = [
        ("CRAFTLIVE_PROJECT_OVERVIEW_JA.md", "全体アーキテクチャ", "有用。ただしCatalog件数・一部素材は現行Asset優先。"),
        ("CRAFTLIVE_INSPECTOR_CONFIGURATION_GUIDE_JA.md", "Inspector全項目", "最も詳細。現行Scene値との差分は本書/Scene優先。"),
        ("PROJECT_STATE_JA.md", "進捗記録", "本文のStep9未実装と追記の完了記録が混在。最終追記を履歴として読む。"),
        ("README_SETUP_JA.md", "V2 setup", "旧方式。消費在庫、Pad1 QR、11素材/3武器は現行でない。"),
        ("STEP0_BASELINE_AUDIT_JA.md", "移行前棚卸し", "履歴。旧Craft Scene構造の理解に使用。"),
        ("STEP1_V3_SETUP_JA.md", "schema/data移行", "V3導入履歴。現行はschema v5。"),
        ("STEP2_FOUR_PAD_SETUP_JA.md", "4 Scene化", "Bootstrap/roleの由来。Transport disabled前提は旧。"),
        ("STEP3_PAD1_SETUP_JA.md", "Pad1構築", "画廊初期実装の履歴。現行Inspectorを優先。"),
        ("STEP4_PAD2_SETUP_JA.md", "Pad2構築", "作業台初期実装の履歴。武器数/Rules差分に注意。"),
        ("STEP56_TRANSFER_PAD3_SETUP_JA.md", "転送/Pad3", "現行フローの由来。Watchdog/登録方式を現行コードで確認。"),
        ("STEP78_SYNTHESIS_SESSION_SETUP_JA.md", "合成/セッション", "ハンマー/履歴/最終選択の設計履歴。"),
        ("STEP9_WEBGL_IPAD_SETUP_JA.md", "WebGL/iPad", "公開設定の基礎。2026-08-08 ProductionReadinessで再確認。"),
    ]
    add_table(doc, ["文書", "主題", "現行での扱い"], rows, [3100, 1900, 4360], font_size=7.6)


def add_scene_appendix(doc, scenes):
    add_page_break(doc)
    doc.add_heading("付録B. Scene階層とCraft-liveコンポーネント", level=1)
    for scene in scenes:
        if "Assets/Scenes/CraftLive/" not in scene["path"]:
            continue
        doc.add_heading(Path(scene["path"]).stem, level=2)
        p = doc.add_paragraph(f"Path: {scene['path']} / GameObjects: {scene['game_object_count']} / Roots: {scene['root_count']}")
        p.style = "Small Note"
        hierarchy = []
        for line in scene["hierarchy"]:
            clean = line.replace("窶・", "—")
            hierarchy.append((clean,))
        add_table(doc, ["Hierarchy（active/inactive・主要component）"], hierarchy, [9360], font_size=6.8, zebra=False)
        if scene["source_prefabs"]:
            rows = [(path, count) for path, count in scene["source_prefabs"].items()]
            add_table(doc, ["Source Prefab", "Instances"], rows, [7860, 1500], font_size=7.0)


def add_inspector_appendix(doc, scenes):
    add_page_break(doc)
    doc.add_heading("付録C. 現行Sceneのserialized Inspector値", level=1)
    add_callout(doc, "読み方", "Unity YAMLからCraft-liveスクリプトのserialized fieldを抽出した値です。fileID/GUID参照はScene/Asset内参照を表し、実行時生成値やprivate非serialized値は含みません。最終判断はUnity Inspectorと実行時ログで確認してください。", "info")
    for scene in scenes:
        if "Assets/Scenes/CraftLive/" not in scene["path"]:
            continue
        doc.add_heading(Path(scene["path"]).stem, level=2)
        rows = []
        for comp in scene["component_settings"]:
            for field, value in comp["fields"].items():
                if len(value) > 140:
                    value = value[:137] + "..."
                rows.append((comp["object"], comp["component"], field, value))
        add_table(doc, ["Object", "Component", "Field", "Serialized value"], rows, [2350, 2200, 2250, 2560], font_size=6.25)


def add_serialized_fields_appendix(doc, scripts):
    add_page_break(doc)
    doc.add_heading("付録D. スクリプト別[SerializeField]一覧", level=1)
    rows = []
    for s in scripts:
        for typ, name in s["fields"]:
            rows.append((s["name"], typ, name, s["path"]))
    add_table(doc, ["Script", "Type", "Field", "Path"], rows, [1900, 1850, 2450, 3160], font_size=6.45)
    p = doc.add_paragraph(f"抽出件数: {len(rows)} fields / {len(scripts)} C# files。UnityEventやpublic field、SerializeReferenceの特殊表現は別途コードを参照。")
    p.style = "Small Note"


def add_manifest_appendix(doc, manifest):
    add_page_break(doc)
    doc.add_heading("付録E. Assets / Packages / ProjectSettings 完全ファイル一覧", level=1)
    counts = Counter(rel.split("/")[0] for rel, _, _, _ in manifest)
    meta_count = sum(1 for _, ext, _, _ in manifest if ext == ".meta")
    add_callout(doc, "完全性の基準", f"以下は生成時点の全{len(manifest)}ファイルです（Assets={counts['Assets']}、Packages={counts['Packages']}、ProjectSettings={counts['ProjectSettings']}、.meta={meta_count}）。引継ぎ後のsnapshotと比較し、特に.meta欠落を検出してください。Library/Logs/Buildsは含みません。", "info")
    rows = [(rel, ext, fmt_size(size), role) for rel, ext, size, role in manifest]
    add_table(doc, ["Path", "Ext", "Size", "Role"], rows, [5150, 750, 1100, 2360], font_size=5.95)


def add_final_checklist(doc):
    add_page_break(doc)
    doc.add_heading("付録F. 引継ぎ完了チェックシート", level=1)
    checks = [
        ("ソース", "Assets/Packages/ProjectSettings/.metaの完全snapshotあり", "□"),
        ("Git", "未追跡0（意図したignoreを除く）、commit hash記録", "□"),
        ("LFS", "新cloneで全binary実体を取得", "□"),
        ("GitHub", "新owner/Adminでpush・branch settings確認", "□"),
        ("Firebase", "新ownerでRules編集・データread/write確認", "□"),
        ("Security", "Rulesがpublic/privateの意図を承認", "□"),
        ("Hosting", "新ownerでdeploy/rollback/cache purge確認", "□"),
        ("DNS/TLS", "更新権限・課金・2FA・期限確認", "□"),
        ("Unity", "6000.4.0f1でcompile 0 error", "□"),
        ("Validation", "CurrentValidation 0 error", "□"),
        ("Production", "ProductionReadiness 0 error", "□"),
        ("Tests", "既知2失敗の処置方針、追加失敗0", "□"),
        ("Build", "現行commitからWebGL再build", "□"),
        ("E2E", "4台でQR→転送→合成→最終code", "□"),
        ("Offline", "切断/復帰とroom分離を確認", "□"),
        ("Operations", "端末・Wi-Fi・問い合わせ・ログ保存先を移管", "□"),
        ("Old account", "新環境の復旧確認後に旧権限を削除", "□"),
    ]
    add_table(doc, ["区分", "確認内容", "完了"], checks, [1600, 6860, 900], font_size=8.1)
    doc.add_paragraph()
    add_table(doc, ["署名", "氏名 / 日付 / 備考"], [
        ("旧所有者", ""),
        ("新所有者", ""),
        ("技術担当", ""),
        ("運用担当", ""),
    ], [2200, 7160], font_size=9.0, zebra=False)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = {
        "architecture": IMAGE_DIR / "architecture.png",
        "flow": IMAGE_DIR / "flow.png",
        "ownership": IMAGE_DIR / "ownership.png",
    }
    make_architecture_diagram(diagrams["architecture"])
    make_flow_diagram(diagrams["flow"])
    make_ownership_diagram(diagrams["ownership"])

    scenes = load_scene_analysis()
    scripts = collect_scripts()
    manifest = collect_manifest()

    doc = Document()
    style_document(doc)
    doc.core_properties.title = "CraftOrigin / Craft-live アカウント引継ぎ・プロジェクト全体詳細書"
    doc.core_properties.subject = "Unity WebGL project handoff reference"
    doc.core_properties.author = "Codex / project workspace audit"
    doc.core_properties.keywords = "CraftOrigin, Craft-live, Unity, WebGL, Firebase, handoff, 引継ぎ"
    doc.core_properties.comments = "Generated from the current workspace on 2026-08-09 JST."

    add_cover(doc)
    add_document_control(doc)
    add_exec_summary(doc)
    add_git_critical(doc)
    add_architecture(doc, diagrams)
    add_flow_and_state(doc, diagrams)
    add_accounts(doc, diagrams)
    add_setup_build(doc)
    add_scenes(doc, scenes)
    add_state_model(doc)
    add_transport(doc)
    add_pads_detail(doc)
    add_catalog(doc)
    add_rules_calc(doc)
    add_web_qr(doc)
    add_project_structure(doc, manifest)
    add_code_inventory(doc, scripts)
    add_quality_risks(doc)
    add_handoff_sop(doc)
    add_operations(doc)
    add_existing_docs(doc)
    add_scene_appendix(doc, scenes)
    add_inspector_appendix(doc, scenes)
    add_serialized_fields_appendix(doc, scripts)
    add_manifest_appendix(doc, manifest)
    add_final_checklist(doc)

    doc.save(OUT_DOCX)
    print(json.dumps({
        "output": str(OUT_DOCX),
        "scripts": len(scripts),
        "manifest": len(manifest),
        "scenes": len(scenes),
        "serialized_fields": sum(len(s["fields"]) for s in scripts),
    }, ensure_ascii=False, indent=2))


def make_qa_doc(title: str) -> Document:
    doc = Document()
    style_document(doc)
    doc.core_properties.title = title
    doc.core_properties.author = "Codex / layout QA split"
    p = doc.add_paragraph()
    p.style = "Title"
    r = p.add_run(title)
    set_run_font(r, size=22, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph("レイアウト検査用の分割版。納品対象は統合DOCX。")
    p.style = "Small Note"
    return doc


def build_qa_parts():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    parts_dir = QA_DIR / "parts"
    parts_dir.mkdir(parents=True, exist_ok=True)
    IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = {
        "architecture": IMAGE_DIR / "architecture.png",
        "flow": IMAGE_DIR / "flow.png",
        "ownership": IMAGE_DIR / "ownership.png",
    }
    make_architecture_diagram(diagrams["architecture"])
    make_flow_diagram(diagrams["flow"])
    make_ownership_diagram(diagrams["ownership"])
    scenes = load_scene_analysis()
    scripts = collect_scripts()
    manifest = collect_manifest()

    docs = []
    main1 = make_qa_doc("QA 1 / 本文：管理・要点・Git・全体像・フロー・アカウント")
    add_document_control(main1)
    add_exec_summary(main1)
    add_git_critical(main1)
    add_architecture(main1, diagrams)
    add_flow_and_state(main1, diagrams)
    add_accounts(main1, diagrams)
    docs.append(("qa_01_main_accounts.docx", main1))

    main2 = make_qa_doc("QA 2 / 本文：環境・Scene・状態・同期")
    add_setup_build(main2)
    add_scenes(main2, scenes)
    add_state_model(main2)
    add_transport(main2)
    docs.append(("qa_02_main_system.docx", main2))

    main3 = make_qa_doc("QA 3 / 本文：Pad・Catalog・Rules・WebGL")
    add_pads_detail(main3)
    add_catalog(main3)
    add_rules_calc(main3)
    add_web_qr(main3)
    docs.append(("qa_03_main_features.docx", main3))

    main4 = make_qa_doc("QA 4 / 本文：構成・コード・品質・SOP・運用")
    add_project_structure(main4, manifest)
    add_code_inventory(main4, scripts)
    add_quality_risks(main4)
    add_handoff_sop(main4)
    add_operations(main4)
    add_existing_docs(main4)
    docs.append(("qa_04_main_operations.docx", main4))

    scene_doc = make_qa_doc("QA 5 / Scene階層")
    add_scene_appendix(scene_doc, scenes)
    docs.append(("qa_05_scenes.docx", scene_doc))

    inspector_doc = make_qa_doc("QA 6 / Inspector値")
    add_inspector_appendix(inspector_doc, scenes)
    docs.append(("qa_06_inspector.docx", inspector_doc))

    fields_doc = make_qa_doc("QA 7 / SerializeField一覧")
    add_serialized_fields_appendix(fields_doc, scripts)
    docs.append(("qa_07_fields.docx", fields_doc))

    manifest_chunks = [
        ("Assets", [row for row in manifest if row[0].startswith("Assets/")]),
        ("Packages_ProjectSettings", [row for row in manifest if not row[0].startswith("Assets/")]),
    ]
    for idx, (label, subset) in enumerate(manifest_chunks, start=8):
        manifest_doc = make_qa_doc(f"QA {idx} / 完全ファイル一覧 {label}")
        add_manifest_appendix(manifest_doc, subset)
        docs.append((f"qa_{idx:02d}_manifest_{label}.docx", manifest_doc))

    checklist_doc = make_qa_doc("QA 10 / 完了チェック")
    add_final_checklist(checklist_doc)
    docs.append(("qa_10_checklist.docx", checklist_doc))

    outputs = []
    for filename, doc in docs:
        path = parts_dir / filename
        doc.save(path)
        outputs.append(str(path))
    print(json.dumps({"qa_parts": outputs}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--qa-parts", action="store_true")
    args = parser.parse_args()
    if args.qa_parts:
        build_qa_parts()
    else:
        build()
